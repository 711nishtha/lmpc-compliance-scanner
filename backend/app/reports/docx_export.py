"""Editable DOCX compliance report export (python-docx) — same content as the PDF, editable
so an enforcement officer can annotate it or attach it to a formal notice."""
from __future__ import annotations

import datetime as dt

from docx import Document
from docx.shared import Pt

from app.rules.schema import ComplianceReport


def generate_docx_report(
    output_path: str,
    report: ComplianceReport,
    product_name: str,
    inspector_email: str | None,
    font_size_tier: str = "tier1",
    image_quality_warning: str | None = None,
) -> None:
    doc = Document()
    doc.add_heading("Legal Metrology Packaged Commodities — Compliance Report", level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Product: {product_name}\n").bold = True
    meta.add_run(f"Inspector: {inspector_email or 'N/A'}\n")
    meta.add_run(f"Scan date: {dt.datetime.utcnow().isoformat(timespec='seconds')} UTC\n")
    meta.add_run(f"Ruleset version: {report.ruleset_version}\n")

    if image_quality_warning:
        warning_p = doc.add_paragraph()
        warning_p.add_run("Possible bad photo, not a bad product: ").bold = True
        warning_p.add_run(image_quality_warning)

    doc.add_heading("Overall status", level=2)
    score = report.compliance_score
    score_text = f"{score}% of applicable checks passed" if score is not None else "N/A"
    doc.add_paragraph(
        f"{report.overall_status.value} — {report.pass_count} PASS, {report.fail_count} FAIL, "
        f"{report.needs_verification_count} NEEDS VERIFICATION ({score_text})."
    )

    doc.add_heading("Itemized rule-cited results", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Rule", "Requirement", "Status", "Evidence / Notes"
    for r in report.results:
        row = table.add_row().cells
        row[0].text = f"{r.rule_id} ({r.rule_reference})"
        row[1].text = r.requirement_text
        row[2].text = r.status.value
        evidence_text = r.evidence.extracted_value or ""
        notes = r.notes or ""
        row[3].text = " / ".join(x for x in (evidence_text, notes) if x) or "—"

    doc.add_heading("Methodology", level=2)
    tier_note = (
        "Tier 2 (calibrated, mm-accurate against Rule 7 tables)"
        if font_size_tier == "tier2"
        else "Tier 1 (relative comparison only — not a calibrated millimetre measurement)"
    )
    p = doc.add_paragraph(
        "Declarations were extracted via offline OCR (Tesseract, eng+hin+guj) followed by "
        "regex/keyword-anchored structured extraction. Low-confidence OCR reads are marked "
        f"NEEDS VERIFICATION rather than auto-passed. Font-size findings used {tier_note}. "
        "Items marked 'VERIFY WITH DoCA' in docs/LEGAL_REQUIREMENTS.md never resolve to an "
        "automated PASS/FAIL. R8-1 (grouped on the principal display panel) is a 2D image-plane "
        "proximity proxy for \"same panel\", not a certified panel determination. R8-2 "
        "(net-quantity clear space) checks Rule 8(1)'s proviso directly but approximates numeral "
        "height from the OCR text line. See LEGAL_REQUIREMENTS.md §10. This report is decision "
        "support for enforcement officials, not a final legal determination."
    )
    for run in p.runs:
        run.font.size = Pt(9)

    doc.save(output_path)
